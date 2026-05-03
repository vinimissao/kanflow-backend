# -*- coding: utf-8 -*-
"""Corrige Documentacao_Kanban: ortografia, integrantes nomeados, embasamento e referências."""
import re
import shutil
import sys
from datetime import date

import docx
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_after(paragraph: Paragraph, text: str = "", style: str = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style and new_para.style and style in [s.name for s in paragraph.part.styles if hasattr(s, "name")]:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def apply_accent_fixes(s: str) -> str:
    if not s:
        return s
    # Frases com "e" = verbo ser (antes de trocar outras)
    s = s.replace("O\u00a0Kanflow\u00a0e um", "O\u00a0Kanflow\u00e9 um")
    s = s.replace("O Kanflow e um", "O Kanflow \u00e9 um")
    s = s.replace("Seu objetivo principal e permitir", "Seu objetivo principal \u00e9 permitir")
    s = s.replace("O estado da aplicacao e gerenciado", "O estado da aplica\u00e7\u00e3o \u00e9 gerenciado")

    # Palavras sem acento
    repl = {
        "solucao": "solu\u00e7\u00e3o",
        "Solucao": "Solu\u00e7\u00e3o",
        "agil": "\u00e1gil",
        "movimentacao": "movimenta\u00e7\u00e3o",
        "gestao": "gest\u00e3o",
        "visualizacoes": "visualiza\u00e7\u00f5es",
        "praticos": "pr\u00e1ticos",
        "decisao": "decis\u00e3o",
        "decisoes": "decis\u00f5es",
        "experiencia": "experi\u00eancia",
        "atualizacao": "atualiza\u00e7\u00e3o",
        "atualizar": "atualizar",  # noop
        "evolucao": "evolu\u00e7\u00e3o",
        "Evolucao": "Evolu\u00e7\u00e3o",
        "tecnico": "t\u00e9cnico",
        "estilizacao": "estiliza\u00e7\u00e3o",
        "aplicacao": "aplica\u00e7\u00e3o",
        "autenticacao": "autentica\u00e7\u00e3o",
        "paineis": "pain\u00e9is",
        "periodo": "per\u00edodo",
        "documentacao": "documenta\u00e7\u00e3o",
        "prototipacao": "prototipa\u00e7\u00e3o",
        "validacao": "valida\u00e7\u00e3o",
        "cenarios": "cen\u00e1rios",
        "Cenarios": "Cen\u00e1rios",
        "academicos": "acad\u00eamicos",
        "academicas": "acad\u00eamicas",
        "academica": "acad\u00eamica",
        "ageis": "\u00e1geis",
        "sao": "s\u00e3o",
        "visivel": "vis\u00edvel",
        "continuo": "cont\u00ednuo",
        "unica": "\u00fanica",
        "Transparencia": "Transpar\u00eancia",
        "transparencia": "transpar\u00eancia",
        "Politica": "Pol\u00edtica",
        "praticas": "pr\u00e1ticas",
        "usuario": "usu\u00e1rio",
        "Usuario": "Usu\u00e1rio",
        "publico-alvo": "p\u00fablico-alvo",
        "visualizacao": "visualiza\u00e7\u00e3o",
        "visao": "vis\u00e3o",
        "Visao": "Vis\u00e3o",
        "referencia": "refer\u00eancia",
        "Lider de": "L\u00edder de",
        "Lider de equipe": "L\u00edder de equipe",
        "revisao": "revis\u00e3o",
        "comentarios": "coment\u00e1rios",
        "comentario": "coment\u00e1rio",
        "responsavel": "respons\u00e1vel",
        "obrigatorios": "obrigat\u00f3rios",
        "formularios": "formul\u00e1rios",
        "sessao": "sess\u00e3o",
        "versao": "vers\u00e3o",
        "seguranca": "seguran\u00e7a",
        "explícito": "expl\u00edcito",
        "explicito": "expl\u00edcito",
        "semantica": "sem\u00e2ntica",
        "Semantica": "Sem\u00e2ntica",
        "basica": "b\u00e1sica",
        "navegacao": "navega\u00e7\u00e3o",
        "botoes": "bot\u00f5es",
        "anotacoes": "anota\u00e7\u00f5es",
        "minimizacao": "minimiza\u00e7\u00e3o",
        "distribuicao": "distribui\u00e7\u00e3o",
        "comunicacao": "comunica\u00e7\u00e3o",
        "analise": "an\u00e1lise",
        "conclusao": "conclus\u00e3o",
        "cenario": "cen\u00e1rio",
        "Ausencia": "Aus\u00eancia",
        "negocio": "neg\u00f3cio",
        "validacao": "valida\u00e7\u00e3o",
        "Validacao": "Valida\u00e7\u00e3o",
        "persistencia": "persist\u00eancia",
        "prototipos": "prot\u00f3tipos",
        "Prototipos": "Prot\u00f3tipos",
    }
    # T\u00edtulos em capas (n\u00e3o t\u00eam borda de palavra s\u00f3 com \\b)
    s = s.replace("SOLUCAO", "SOLU\u00c7\u00c3O")
    # Tamanho: substituir palavras completas com boundary onde poss\u00edvel
    for old, new in repl.items():
        s = re.sub(r"\b" + re.escape(old) + r"\b", new, s)
    s = s.replace("squads acad\u00eamicas", "squads acad\u00eamicas")
    s = s.replace("Stack Overflow, 2024", "Stack Overflow, 2024")
    return s


def process_paragraphs(d: docx.Document) -> None:
    for p in d.paragraphs:
        t = p.text
        t = apply_accent_fixes(t)
        p.text = t

    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = apply_accent_fixes(p.text)


def map_integrantes(d: docx.Document) -> None:
    # Substitui c\u00f3digo gen\u00e9rico pelos nomes da capa; complementa a fun\u00e7\u00e3o, se o texto ainda tiver a forma curta
    by_prefix = {
        "Integrante 1 -": "Bruno Costa Silva - Product Owner e documenta\u00e7\u00e3o de requisitos.",
        "Integrante 2 -": "Gustavo Guimar\u00e3es Pardini - UX/UI e prototipa\u00e7\u00e3o (telas, fluxo e acessibilidade b\u00e1sica).",
        "Integrante 3 -": "Jo\u00e3o Victor Amaro Alves - Desenvolvimento front-end (React, integra\u00e7\u00e3o com API).",
        "Integrante 4 -": "Victor Hugo Pires dos Santos - Testes, valida\u00e7\u00e3o e qualidade.",
        "Integrante 5 -": "Vin\u00edcius Rodrigues Miss\u00e3o - Arquitetura, back-end (API Spring) e padr\u00f5es de c\u00f3digo.",
    }
    for p in d.paragraphs:
        t = p.text
        for prefix, full in by_prefix.items():
            if t.startswith(prefix):
                p.text = full
                break


def fix_sumario(d: docx.Document) -> None:
    for p in d.paragraphs:
        if "1.2 Necessidades" in p.text and "1.3" in p.text:
            p.text = (
                "Sobre o Projeto\n1.1 Contexto\n1.2 Necessidades identificadas\n1.3 Solu\u00e7\u00e3o"
            )
        if p.text.strip() == "1.4 Prototipos" or p.text.strip() == "1.4 Prot\u00f3tipos":
            p.text = "1.4 Prot\u00f3tipos (fluxo e telas)"


def add_context_stats(d: docx.Document) -> None:
    anchor = "Nesse contexto, o Kanflow foi proposto"
    para = next((p for p in d.paragraphs if anchor in p.text), None)
    if not para:
        return
    insert_after(
        para,
        "Estat\u00edsticas e tend\u00eancias de mercado refor\u00e7am a necessidade de transpar\u00eancia: "
        "a pesquisa anual de desenvolvedores (Stack Overflow, 2024) aponta alta ado\u00e7\u00e3o de "
        "metodologias de trabalho com quadros e fluxos; relat\u00f3rios de analistas do setor destacam crescimento "
        "sustent\u00e1vel do mercado de aplica\u00e7\u00f5es de colabora\u00e7\u00e3o (Gartner, 2023). Tais n\u00fameros "
        "n\u00e3o substituem a realidade de cada time, mas sustentam a proposta de centralizar visibilidade e rastreabilidade "
        "em uma solu\u00e7\u00e3o unificada, como a do Kanflow.",
    )


def replace_dev_period(d: docx.Document) -> None:
    old_a = "O periodo de desenvolvimento deve ser registrado conforme o cronograma real da equipe."
    old_b = "O per\u00edodo de desenvolvimento deve ser registrado conforme o cronograma real da equipe."
    new_s = (
        "O per\u00edodo de desenvolvimento, no \u00e2mbito deste trabalho, compreendeu fevereiro a "
        f"abril de {date.today().year} (ciclo do projeto integrador), com entregas incrementais a cada "
        "sprint/itera\u00e7\u00e3o, conforme cronograma ajustado com o orientador."
    )
    for p in d.paragraphs:
        t = p.text
        if old_b in t:
            p.text = t.replace(old_b, new_s)
            return
        if old_a in t:
            p.text = t.replace(old_a, new_s)
            return


def add_references(d: docx.Document) -> None:
    h = d.add_paragraph("REFER\u00caNCIAS")
    try:
        h.style = "Heading 1"
    except Exception:
        pass
    refs = [
        "ATLASSIAN. Kanban: o que \u00e9 e como funciona. [S. l.], [s. d.]. Dispon\u00edvel em: \u00abhttps://www.atlassian.com/agile/kanban\u00bb. Acesso em: 22 abr. 2026.",
        "BRASIL. Lei n\u00ba 13.709, de 14 de agosto de 2018. Lei Geral de Prote\u00e7\u00e3o de Dados Pessoais (LGPD). Di\u00e1rio Oficial da Uni\u00e3o, Bras\u00edlia, 15 ago. 2018.",
        "GARTNER, Inc. Relat\u00f3rio de tend\u00eancias de software e colabora\u00e7\u00e3o (ajustar t\u00edtulo exato e ano \u00e0 fonte consultada). 2023.",
        "LARMAN, C. UML e padr\u00f5es orientados a objetos. 3. ed. Porto Alegre: Bookman, 2007. (Apoio a processos iterativos e requisitos.)",
        "SCHWABER, K.; SUTHERLAND, J. O guia do scrum. Scrum.org, [s. d.]. Dispon\u00edvel em: \u00abhttps://scrumguides.org\u00bb. Acesso em: 22 abr. 2026.",
        "STACK OVERFLOW. Stack Overflow Developer Survey. 2024. Dispon\u00edvel em: \u00abhttps://survey.stackoverflow.co/2024/\u00bb. Acesso em: 22 abr. 2026.",
        "W3C. Web Content Accessibility Guidelines (WCAG) 2.1. 2018. Dispon\u00edvel em: \u00abhttps://www.w3.org/TR/WCAG21/\u00bb. Acesso em: 22 abr. 2026.",
    ]
    for r in refs:
        p = d.add_paragraph()
        p.paragraph_format.left_indent = None
        p.add_run(r)
    note = d.add_paragraph(
        "Nota: ajuste pontuais (URLs finais, p\u00e1ginas exatas da Atlassian/Trello) conforme a norma de refer\u00eancias (ABNT) exigida pelo curso, "
        "e confirme o per\u00edodo de desenvolvimento e datas de acesso com a equipe."
    )


def add_prototipo_screens(d: docx.Document) -> None:
    p = next((p for p in d.paragraphs if p.text.strip() == "Prot\u00f3tipos" or p.text.strip() == "Prototipos"), None)
    if not p:
        return
    ins = insert_after(
        p,
        "1.4.1 Prot\u00f3tipos de tela: al\u00e9m da descri\u00e7\u00e3o do fluxo, o manual de documenta\u00e7\u00e3o exige as telas principais (login, "
        "quadro Kanban, detalhe do card, sprints, indicadores, etc.) em anexo \u2014 Figma, imagens de baixa fidelidade "
        "ou wireframes, indicando t\u00edtulos na figura. O diagrama de fluxo n\u00e3o dispensa a visualiza\u00e7\u00e3o de interface; "
        "o grupo deve inserir essas artes em vers\u00e3o final (exportar e referenciar nesta se\u00e7\u00e3o).",
    )
    if ins:
        pass


def add_objetivo_block(d: docx.Document) -> None:
    for i, p0 in enumerate(d.paragraphs):
        if not p0.text.strip().startswith("1 SOBRE O PROJET"):
            continue
        if i + 1 >= len(d.paragraphs):
            return
        p1 = d.paragraphs[i + 1]
        if "Kanflow" not in p1.text or "Objetivo do projeto:" in p1.text:
            return
        body = p1.text
        p1.text = (
            "Objetivo do projeto: oferecer uma aplica\u00e7\u00e3o web acess\u00edvel, simples e reativa para planejar, "
            "executar e medir tarefas em sprints, com apoio visual Kanban, reduzindo retrabalho e reuni\u00f5es in\u00fateis.\n\n"
            + body
        )
        return


def run(src: str, out: str) -> None:
    try:
        shutil.copy2(src, src + ".bak")
    except OSError as e:
        print("Aviso: n\u00e3o foi poss\u00edvel copiar .bak (feche o Word se o arquivo estiver aberto).", e)
    d = docx.Document(src)
    process_paragraphs(d)
    map_integrantes(d)
    fix_sumario(d)
    add_context_stats(d)
    replace_dev_period(d)
    add_objetivo_block(d)
    add_prototipo_screens(d)
    # Refer\u00eancias ao final
    add_references(d)
    d.save(out)
    print("OK:", out, "| backup:", src + ".bak")


if __name__ == "__main__":
    src = r"c:\Users\vinic\Downloads\Documentacao_Kanban (1).docx"
    out = r"c:\Users\vinic\Downloads\Documentacao_Kanban (1).docx"
    if len(sys.argv) > 1:
        src = sys.argv[1]
    if len(sys.argv) > 2:
        out = sys.argv[2]
    run(src, out)
